from __future__ import annotations

import asyncio
import io
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.logging import get_logger
from app.models.document import Document, DocumentChunk
from app.services.memory_service import _get_embedding_model

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".docx"}

# Hard ceiling so a single upload cannot occupy the event loop indefinitely
MAX_CHUNKS_PER_DOCUMENT = 1000

# Batch size for Qdrant upserts
_UPSERT_BATCH_SIZE = 100

_qdrant_client = None


def _get_qdrant():
    global _qdrant_client
    if _qdrant_client is None:
        from qdrant_client import QdrantClient
        _qdrant_client = QdrantClient(
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY,
        )
    return _qdrant_client


# ── Embedding (same approach as memory_service) ───────────────────────────────

def _embed(text: str) -> List[float]:
    model = _get_embedding_model()
    return model.encode(text, normalize_embeddings=True).tolist()


async def _ensure_collection() -> None:
    from qdrant_client.http.models import Distance, VectorParams
    client = _get_qdrant()
    loop = asyncio.get_event_loop()
    collections = await loop.run_in_executor(None, lambda: client.get_collections().collections)
    names = [c.name for c in collections]
    if settings.QDRANT_COLLECTION_DOCS not in names:
        await loop.run_in_executor(
            None,
            lambda: client.create_collection(
                collection_name=settings.QDRANT_COLLECTION_DOCS,
                vectors_config=VectorParams(
                    size=settings.VECTOR_DIMENSION, distance=Distance.COSINE
                ),
            ),
        )


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text(filename: str, content_bytes: bytes) -> str:
    """Extract plain text from an uploaded file. Raises ValueError for unsupported types."""
    suffix = Path(filename or "").suffix.lower()

    if suffix == ".pdf":
        import pdfplumber
        pages: List[str] = []
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)

    if suffix in (".txt", ".md"):
        return content_bytes.decode("utf-8", errors="replace")

    if suffix == ".csv":
        # Keep CSV as raw text — rows stay on their own lines
        return content_bytes.decode("utf-8", errors="replace")

    if suffix == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(content_bytes))
        parts: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    raise ValueError(
        f"Unsupported file type '{suffix or 'unknown'}'. "
        f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


# ── Chunking ──────────────────────────────────────────────────────────────────

def _split_long_paragraph(paragraph: str, chunk_size: int, overlap: int) -> List[str]:
    """Hard-split a paragraph longer than chunk_size, preferring word boundaries."""
    pieces: List[str] = []
    start = 0
    length = len(paragraph)
    while start < length:
        end = min(start + chunk_size, length)
        if end < length:
            space = paragraph.rfind(" ", start, end)
            if space > start:
                end = space
        piece = paragraph[start:end].strip()
        if piece:
            pieces.append(piece)
        if end >= length:
            break
        start = max(end - overlap, start + 1)
    return pieces


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into ~chunk_size character chunks on paragraph boundaries where possible,
    carrying ~overlap characters of trailing context into the next chunk."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    # Normalise into pieces no longer than chunk_size
    pieces: List[str] = []
    for para in paragraphs:
        if len(para) <= chunk_size:
            pieces.append(para)
        else:
            pieces.extend(_split_long_paragraph(para, chunk_size, overlap))

    chunks: List[str] = []
    current: List[str] = []
    current_len = 0
    for piece in pieces:
        if current and current_len + len(piece) + 2 > chunk_size:
            chunks.append("\n\n".join(current))
            # Carry trailing pieces (up to `overlap` chars) into the next chunk
            carried: List[str] = []
            carried_len = 0
            for prev in reversed(current):
                if carried_len + len(prev) > overlap:
                    break
                carried.insert(0, prev)
                carried_len += len(prev) + 2
            current = carried
            current_len = carried_len
        current.append(piece)
        current_len += len(piece) + 2
    if current:
        chunks.append("\n\n".join(current))
    return chunks


# ── Processing pipeline ───────────────────────────────────────────────────────

async def process_document(db: AsyncSession, document: Document, content_bytes: bytes) -> Document:
    """Extract → chunk → embed → upsert into Qdrant → persist chunks.
    Sets document.status to 'ready' or 'failed' (with error message)."""
    loop = asyncio.get_event_loop()
    try:
        text = await loop.run_in_executor(None, extract_text, document.filename, content_bytes)
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("No text content could be extracted from the document")
        if len(chunks) > MAX_CHUNKS_PER_DOCUMENT:
            raise ValueError(
                f"Document produced {len(chunks)} chunks (max {MAX_CHUNKS_PER_DOCUMENT}); "
                "please upload a smaller document"
            )

        await _ensure_collection()
        client = _get_qdrant()
        from qdrant_client.http.models import PointStruct

        points: List[PointStruct] = []
        chunk_rows: List[DocumentChunk] = []
        for idx, chunk in enumerate(chunks):
            vector = await loop.run_in_executor(None, _embed, chunk)
            embedding_id = str(uuid.uuid4())
            points.append(
                PointStruct(
                    id=embedding_id,
                    vector=vector,
                    payload={
                        "user_id": document.user_id,
                        "document_id": str(document.id),
                        "filename": document.filename,
                        "chunk_index": idx,
                        "content": chunk,
                    },
                )
            )
            chunk_rows.append(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx,
                    content=chunk,
                    embedding_id=embedding_id,
                )
            )

        for i in range(0, len(points), _UPSERT_BATCH_SIZE):
            batch = points[i : i + _UPSERT_BATCH_SIZE]
            await loop.run_in_executor(
                None,
                lambda b=batch: client.upsert(
                    collection_name=settings.QDRANT_COLLECTION_DOCS, points=b
                ),
            )

        db.add_all(chunk_rows)
        document.chunk_count = len(chunks)
        document.status = "ready"
        document.error = None
        logger.info(
            "document_processed",
            document_id=str(document.id),
            filename=document.filename,
            chunks=len(chunks),
        )
    except Exception as exc:
        logger.error(
            "document_processing_failed", document_id=str(document.id), error=str(exc)
        )
        document.status = "failed"
        document.error = str(exc)[:2000]

    await db.flush()
    return document


# ── Search ────────────────────────────────────────────────────────────────────

async def search_documents(
    query: str,
    user_id: int,
    document_ids: Optional[List[str]] = None,
    limit: int = 5,
) -> List[Dict[str, Any]]:
    """Semantic search over the user's indexed document chunks.
    Returns a list of {content, document_id, filename, score}."""
    await _ensure_collection()
    loop = asyncio.get_event_loop()
    vector = await loop.run_in_executor(None, _embed, query)
    client = _get_qdrant()

    must: List[Dict[str, Any]] = [{"key": "user_id", "match": {"value": user_id}}]
    if document_ids:
        must.append({"key": "document_id", "match": {"any": [str(d) for d in document_ids]}})

    results = await loop.run_in_executor(
        None,
        lambda: client.search(
            collection_name=settings.QDRANT_COLLECTION_DOCS,
            query_vector=vector,
            limit=limit,
            query_filter={"must": must},
        ),
    )

    return [
        {
            "content": (r.payload or {}).get("content", ""),
            "document_id": (r.payload or {}).get("document_id"),
            "filename": (r.payload or {}).get("filename"),
            "score": r.score,
        }
        for r in results
    ]


# ── CRUD helpers ──────────────────────────────────────────────────────────────

async def get_document(db: AsyncSession, document_id: uuid.UUID, user_id: int) -> Optional[Document]:
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.user_id == user_id)
    )
    return result.scalar_one_or_none()


async def list_documents(
    db: AsyncSession, user_id: int, limit: int = 50, offset: int = 0
) -> List[Document]:
    result = await db.execute(
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    return list(result.scalars().all())


async def delete_document(db: AsyncSession, document: Document) -> None:
    """Delete a document, its chunks, and its Qdrant points."""
    embedding_ids = [c.embedding_id for c in document.chunks if c.embedding_id]
    if embedding_ids:
        try:
            from qdrant_client.http.models import PointIdsList
            loop = asyncio.get_event_loop()
            client = _get_qdrant()
            await loop.run_in_executor(
                None,
                lambda: client.delete(
                    collection_name=settings.QDRANT_COLLECTION_DOCS,
                    points_selector=PointIdsList(points=embedding_ids),
                ),
            )
        except Exception as exc:
            logger.warning("qdrant_delete_failed", error=str(exc))
    await db.delete(document)
    await db.flush()
